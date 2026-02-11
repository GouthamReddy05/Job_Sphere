import io
from PyPDF2 import PdfReader
import docx
from huggingface_hub import InferenceClient
import os
from dotenv import load_dotenv
import re
from langchain.chat_models import init_chat_model
from langchain_core.messages import SystemMessage, HumanMessage


load_dotenv()

# hf_api_key = os.getenv('huggingface_api_key')
# if hf_api_key:
#     client = InferenceClient(token=hf_api_key)
# else:
#     client = InferenceClient() 

os.environ["GEMINI_API_KEY"] = os.getenv("gemini_api_key")

model = init_chat_model("google_genai:gemini-2.5-flash")


def extract_json(text):
    if not text:
        return ""

    text = re.sub(r"```json|```", "", text).strip()

    match = re.search(r"(\{.*\}|\[.*\])", text, re.DOTALL)
    if match:
        return match.group(1)

    return text



def extract_text_from_file(file_source):
    """
    Supports:
    - FastAPI UploadFile
    - File-like object
    - File path (str)
    """

    text = ""

    # --- Case 1: FastAPI UploadFile ---
    if hasattr(file_source, "file") and hasattr(file_source, "filename"):
        filename = file_source.filename
        content_stream = io.BytesIO(file_source.file.read())
        file_source.file.seek(0)

    # --- Case 2: File-like object ---
    elif hasattr(file_source, "read"):
        filename = getattr(file_source, "filename", "")
        content_stream = io.BytesIO(file_source.read())
        file_source.seek(0)

    # --- Case 3: File path ---
    elif isinstance(file_source, str):
        if not os.path.exists(file_source):
            raise FileNotFoundError(f"No such file: '{file_source}'")
        filename = file_source
        content_stream = file_source

    else:
        raise TypeError("Invalid file source provided.")

    # --- Handle PDF ---
    if filename.lower().endswith(".pdf"):
        reader = PdfReader(content_stream)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    # --- Handle DOCX ---
    elif filename.lower().endswith(".docx"):
        document = docx.Document(content_stream)
        for para in document.paragraphs:
            text += para.text + "\n"

    else:
        raise ValueError("Unsupported file type. Please upload PDF or DOCX.")

    return text


# def extract_text_from_file(file_source):
#     """
#     Reads a document from either a file path or a file-like object.

#     Args:
#         file_source: A file path (str) or a file-like object.

#     Returns:
#         The extracted text as a string.
#     """
#     text = ""
#     filename = ""

#     if hasattr(file_source, 'read'):
#         # It's a file-like object (e.g., uploaded file in Flask)
#         filename = getattr(file_source, 'filename', 'uploaded_file')
#         content_stream = io.BytesIO(file_source.read())
#         file_source.seek(0)

#     elif isinstance(file_source, str):
#         # It's a file path
#         if not os.path.exists(file_source):
#             raise FileNotFoundError(f"No such file: '{file_source}'")
#         filename = file_source
#         content_stream = file_source  # Keep path for PyPDF2/docx

#     else:
#         raise TypeError("Input must be a file path (string) or a file-like object.")

#     # --- Handle PDF ---
#     if filename.lower().endswith('.pdf'):
#         if isinstance(content_stream, str):  # file path
#             reader = PdfReader(content_stream)
#         else:  # file-like
#             reader = PdfReader(content_stream)
#         for page in reader.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"

#     # --- Handle DOCX ---
#     elif filename.lower().endswith('.docx'):
#         if isinstance(content_stream, str):  # file path
#             document = docx.Document(content_stream)
#         else:  # file-like
#             document = docx.Document(content_stream)
#         for para in document.paragraphs:
#             text += para.text + "\n"

#     else:
#         raise ValueError("Unsupported file type. Please provide a PDF or DOCX file.")

#     return text




# print(extract_ordered_text_pdf("Resume.pdf"))

def send_text_to_llm(text):
    
    prompt = f"""
    You are an AI that organizes resume text into structured sections.

    Task:
    - Organize the resume text into these exact sections:
      - Name and Contact Information
      - Introduction/Summary
      - Experience
      - Projects
      - Education
      - Skills
      - Certifications
    - If a section uses a different title (e.g., "Profile", "Career Overview"), map it to the most relevant one.
    - If any section is missing, still include it with an empty string ("").

    Input Resume Text:
    {text}

    Output:
    Return ONLY valid JSON in this format:
    {{
      "Name and Contact Information": "",
      "Introduction/Summary": "",
      "Experience": "",
      "Projects": "",
      "Education": "",
      "Skills": "",
      "Certifications": ""
    }}
    """


    # try:
    #     response = client.chat_completion(model="meta-llama/Llama-3.3-70B-Instruct", messages=[
    #         {"role": "system", "content": "You are a helpful assistant."},
    #         {"role": "user", "content": prompt}
    #     ])

    #     return f"{response.choices[0].message.content}"


    # except Exception as e:
    #     return f"⚠️ Error generating response: {str(e)}"

    try:
        response = model.invoke([
            SystemMessage(content="You are a helpful assistant."),
            HumanMessage(content=prompt)
        ])

        return response.content

    except Exception as e:
        return f"⚠️ Error generating response: {str(e)}"
    

def retrieve_skills(text):
    prompt = f"""
Extract all skills mentioned in the resume.

Return ONLY valid JSON array format like this:

["Python", "Machine Learning", "SQL", "TensorFlow"]

Resume:
{text}
"""

    try:
        response = model.invoke([
            HumanMessage(content=prompt)
        ])

        return response.content

    except Exception as e:
        return f"⚠️ Error generating response: {str(e)}"

    
    

def generate_missing_skills(role, candidate_skills):

    prompt = f"""
You are a strict skill gap analyzer.

Target Role: {role}

Candidate Skills (Python list):
{candidate_skills}

Steps:
1. Identify standard required skills for this role.
2. Compare with candidate skills.
3. Return ONLY missing skills.

Output STRICT JSON:

{{
  "Core Technical Skills": [],
  "Programming Languages/Frameworks": [],
  "Tools & Platforms": []
}}

Return ONLY JSON.
"""

    try:
        response = model.invoke([
            HumanMessage(content=prompt)
        ])
        return response.content

    except Exception as e:
        return f"⚠️ Error generating response: {str(e)}"




# text = extract_text_from_file("Resume.pdf")
# print(text)

# structured_text = send_text_to_llm(text)
# print(structured_text)

# skills = retrieve_skills(structured_text)
# print(skills)

# missing_skills = generate_missing_skills("Machine Learning engineer", skills)
# print(missing_skills)
